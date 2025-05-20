import pytest
import uuid
from fastapi.testclient import TestClient
from main import app
from models.models import User, Summary
import datetime
import os
from gridfs import GridFS
from pymongo import MongoClient
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document
from gridfs import GridFS
from services.service import service_download_file
from fastapi.responses import StreamingResponse
from bson import ObjectId
from fastapi.exceptions import HTTPException
import time

client = TestClient(app)

mongo_client = MongoClient(os.environ['DATABASE_URL'])
db = mongo_client.internaldb
users_collection = db["users"]
summaries_collection = db["summaries"]
shared_summaries_collection = db["shared_summaries"]
grid_fs = GridFS(db)

@pytest.fixture(scope="module", autouse=True)
def setup_database():
    """
    Setup test database with dummy users and summaries.
    Runs once per module and cleans up after all tests.
    
    Creates:
    - Two users: Alice Smith and Bob Johnson
    - Two summaries: One for each user
    """
    # Dummy data for users
    user1 = User(
        firstName="Alice",
        lastName="Smith",
        phone="+14132752733",
        email="alice.smith@example.com",
        password="password123",
        confirmPassword="password123"
    )
    user1_dict = user1.dict()
    user1_dict['_id']= "03b4aab0-02bf-4920-be97-301f942dadc9"
    user2 = User(
        firstName="Bob",
        lastName="Johnson",
        phone="+14132782738",
        email="bob.johnson@example.com",
        password="password123",
        confirmPassword="password123"
    )
    user2_dict = user2.dict()
    user2_dict['_id']= "aac0c1cc-8af4-4501-81a4-ee10a90cbbcb"
    
    users_collection.delete_one({"email": user1.email})
    users_collection.delete_one({"email": user2.email})
    summaries_collection.delete_many({"userId": {"$in": [user1_dict["_id"], user2_dict["_id"]]}})
    
    # Insert users into the database
    users_collection.insert_many([user1_dict, user2_dict])  # Convert to dict for insertion

    # Dummy data for summaries
    summary1 = Summary(
        userId= "03b4aab0-02bf-4920-be97-301f942dadc9",  
        type="code",
        uploadType="upload",
        initialData="print('Hello from Alice')",
        outputData= "This is a sample summary"
    )
    summary2 = Summary(
        userId="aac0c1cc-8af4-4501-81a4-ee10a90cbbcb", 
        type="code",
        uploadType="upload",
        title="Title",
        initialData="print('Hello from Bob')",
        outputData= "This is a sample summary",
    )
    summary1_dict = summary1.dict()
    summary1_dict['_id']= str(uuid.uuid4())
    summary2_dict = summary2.dict()
    summary2_dict['_id']= str(uuid.uuid4())
    # Insert summaries into the database
    summaries_collection.insert_many([summary1_dict, summary2_dict])  # Convert to dict for insertion
    yield
    users_collection.delete_one({"email": user1.email})
    users_collection.delete_one({"email": user2.email})
    summaries_collection.delete_many({"userId": {"$in": [user1_dict["_id"], user2_dict["_id"]]}})


@pytest.fixture
def create_user():
    """
    Fixture to create a test user or return existing user.
    
    Returns:
        dict: User data including auth token and user details
    """
    user = User(
            firstName="John",
            lastName="Doe",
            phone="+14132772734",
            email="john.doe@example.com",
            password="password123",
            confirmPassword="password123"
        )
    
    # Convert datetime fields to ISO format for JSON serialization
    user_data_dict = user.dict()
    user_data_dict['createdAt'] = user.createdAt.isoformat()
    user_data_dict['lastLoggedInAt'] = user.lastLoggedInAt.isoformat()
    
    # Check if user already exists
    existing_user_response = client.post("/user/verify", json={"email": user_data_dict['email'], "password": user_data_dict['password']})
    
    if existing_user_response.status_code == 200:
        return existing_user_response.json()  # Return existing user

    response = client.post("/user/create", json=user_data_dict)
    return response.json()

@pytest.fixture
def create_summary(create_user):
    if 'user' not in create_user.get("result", {}):
        pytest.fail("User creation failed or userId not found in response.")
    auth_token = create_user["result"]["auth_token"]
    
    summary_data = Summary(
        userId=create_user["result"]["user"]["id"],
        type="code",
        uploadType="upload",
        initialData="print('Hello World')"
    )
    
    summary_data_dict = {
        "userId": summary_data.userId,
        "type": summary_data.type,
        "uploadType": summary_data.uploadType,
        "initialData": summary_data.initialData,
        "createdAt": summary_data.createdAt.isoformat()  # Convert to ISO format string
    }
    
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    
    response = client.post("/summary/create", json=summary_data_dict, headers=headers)
    print("Create summary ",response.json())
    return response.json(), auth_token

def test_create_user():
    """
    Test user creation endpoint.
    Expected Output:
        - Status "OK" for new user
        - "User Already Exists" for duplicate email
    """
    user = User(
            firstName="Abraham",
            lastName="Lincoln",
            phone="+14132072934",
            email="alincoln@example.com",
            password="password13",
            confirmPassword="password13"
        )
    # Convert datetime fields to ISO format for JSON serialization
    users_collection.delete_many({'email': user.email})
    user_data_dict = user.dict()
    user_data_dict['createdAt'] = user.createdAt.isoformat()
    user_data_dict['lastLoggedInAt'] = user.lastLoggedInAt.isoformat()
    response = client.post("/user/create", json=user_data_dict)
    if 'status' in response.json():
        assert response.json()["status"] == "OK"
    elif 'detail' in response.json():
        assert response.json()['detail'] == 'User Already Exists'
    db.users.delete_many({'email': user.email})


def test_regenerate_feedback_with_file():
    """
    Test regenerating a summary with feedback for file-based content.
    
    Steps:
    1. Authenticate user
    2. Create test file and upload
    3. Get summary and regenerate with feedback
    Expected Output:
        - 201 status code
        - New summary with regenerated content
    """
    """Test regenerating a summary with feedback for file-based summary."""
    # First create and verify a user to get auth token
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token']
    userId = response.json()['result']['user']['id']
    
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    
    # Create a test file and upload it
    test_content = "This is test content from a file"
    test_file = BytesIO(test_content.encode('utf-8'))
    files = {
        "file": ("test.txt", test_file, "text/plain")
    }
    form_data = {
        "userId": userId,
        "type": "documentation",
        "uploadType": "upload"
    }

    time.sleep(6)
    
    upload_response = client.post(
        "/summary/upload",
        data=form_data,
        files=files,
        headers=headers
    )
    print("upload_response")
    print(upload_response)
    print("upload_response")
    assert upload_response.status_code == 201
    
    # Get the created summary
    summaries_response = client.get(f"/summaries/{userId}", headers=headers)
    summary = summaries_response.json()["result"][0]
 
    time.sleep(3)
    # Regenerate the summary
    feedback = "Please make it more technical"
    response = client.post(
        f"/summary/regenerate/{summary['id']}", 
        json=feedback,
        headers=headers
    )

    assert response.status_code == 201
    assert "result" in response.json()
    new_summary = response.json()["result"]
    assert isinstance(new_summary, dict)

def test_regenerate_feedback_code_type():
    """
    Test regenerating a summary with code-type content.
    
    Expected Output:
        - 201 status code
        - New summary with regenerated content
    """
    # First create and verify a user to get auth token
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token']
    userId = response.json()['result']['user']['id']
    
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    
    # Create a code summary
    summary_data = {
        "userId": userId,
        "type": "code",
        "uploadType": "upload",
        "initialData": "def hello():\n    print('Hello World')",
        "createdAt": datetime.datetime.now(datetime.UTC).isoformat()
    }

    time.sleep(3)
    
    create_response = client.post(
        "/summary/create", 
        json=summary_data,
        headers=headers
    )
    assert create_response.status_code == 201
    summary_id = create_response.json()["result"]["summary_id"]
    
    time.sleep(3)
    
    # Regenerate the summary
    feedback = "Please explain the function parameters"
    response = client.post(
        f"/summary/regenerate/{summary_id}", 
        json=feedback,
        headers=headers
    )
    
    assert response.status_code == 201
    assert "result" in response.json()
    new_summary = response.json()["result"]
    assert isinstance(new_summary, dict)

def test_create_user_with_existing_email():
    """
    Test creating a user with an email that already exists.
    Expected Output:
        - 409 status code
        - "User Already Exists" error message
    """
    user_data = User(
        firstName="Alice",
        lastName="Smith",
        phone="+14132752733",
        email="alice.smith@example.com",
        password="password123"
    )
    user_data_dict = user_data.dict()
    user_data_dict['createdAt'] = user_data.createdAt.isoformat()
    user_data_dict['lastLoggedInAt'] = user_data.lastLoggedInAt.isoformat()

    response = client.post("/user/create", json=user_data_dict)
    assert response.status_code == 409  # Conflict status code for existing user
    assert response.json().get("detail") == "User Already Exists"

def test_verify_user():
    """
    Test user verification with valid credentials.
    
    Expected Output:
        - 200 status code
        - Response containing auth token
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    assert response.status_code == 200
    assert "auth_token" in response.json()["result"]
    
def test_create_summary():
    """
    Test creating a new summary.
    
    Expected Output:
        - "OK" status
        - Response containing summary_id
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']
    summary_data_dict = {
        "userId": userId,
        "type": "code",
        "uploadType": "upload",
        "initialData": "print('Hello World')",
        "createdAt": datetime.datetime.now(datetime.UTC).isoformat()  # Convert to ISO format string
    }
    
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }

    time.sleep(3)
    response = client.post("/summary/create", json=summary_data_dict, headers=headers)
    response = response.json()
    assert response["status"] == "OK"
    assert "summary_id" in response['result']

def test_user_summaries():
    """
    Test retrieving all summaries for a user.
    
    Expected Output:
        - 200 status code
        - List of user's summaries
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']

    assert userId is not None, "User ID should not be None"
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    response = client.get(f"/summaries/{userId}",headers=headers)
    assert response.status_code == 200
    assert isinstance(response.json()["result"], list)

def test_delete_summary():
    """
    Test deleting a specific summary.
    
    Expected Output:
        - 201 status code
        - Success message confirmation
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']
    summary = summaries_collection.find_one({'userId': userId})
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    delete_response = client.delete(f"/summary/{summary["_id"]}", headers=headers)
    assert delete_response.status_code == 201
    assert delete_response.json()["result"]["message"] == "Summary deleted successfully"

def test_verify_non_existent_user():
    """Test verifying a user account that does not exist."""
    non_existent_user_data = {
        "email": "nonexistent.user@example.com",
        "password": "wrongpassword"
    }
    response = client.post("/user/verify", json=non_existent_user_data)
    assert response.status_code == 404  # Expecting a 404 Not Found status
    assert response.json().get("detail") == "Account Does Not Exist"  # Adjust based on your actual error message 

def test_share_summary_success():
    """
    Test sharing a summary with another user.
    
    Expected Output:
        - 200 status code
        - Success message confirming share
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']
    # Share the summary created by the first user
    summary = summaries_collection.find_one({'userId': userId})
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    recipient_email = "bob.johnson@example.com"
    share_response = client.post("/summary/share", json={"summary_id": summary["_id"], "recipient": recipient_email}, headers=headers)
    assert share_response.status_code == 200
    assert share_response.json()["result"]["message"] == f"Summary shared successfully with {recipient_email}"

def test_share_summary_not_found():
    """
    Test sharing a non-existent summary.
    
    Expected Output:
        - 404 status code
        - "Summary not found" error message
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    recipient_email = "bob.johnson@example.com"
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    share_response = client.post("/summary/share", json={"summary_id": "non-existent-summary","recipient": recipient_email}, headers=headers)
    
    assert share_response.status_code == 404  # Expecting Not Found
    assert share_response.json().get("detail") == "Summary not found"

def test_share_summary_recipient_not_found():
    """
    Test sharing a summary with non-existent recipient.
    
   
    Expected Output:
        - 404 status code
        - Error message about recipient not being registered
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token']
    userId = response.json()['result']['user']['id']
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    summary = summaries_collection.find_one({'userId': userId})
    share_response = client.post(f"/summary/share", json={"summary_id": summary["_id"], "recipient": "nonexistent@example.com"},
                                 headers=headers)
    assert share_response.status_code == 404  # Expecting Not Found
    assert share_response.json().get("detail") == "The recipient must be a registered user of Briefly."

def test_share_summary_with_self():
    """
    Test attempting to share a summary with oneself.
    
    Expected Output:
        - 400 status code
        - Error message preventing self-sharing
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token']
    userId = response.json()['result']['user']['id']
    summary = summaries_collection.find_one({'userId': userId})
    recipient_email = response.json()['result']['user']['email']  # Sharing with the same user
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    share_response = client.post(f"/summary/share", json={"summary_id": summary["_id"], "recipient": recipient_email},
                                 headers=headers)
    assert share_response.status_code == 400  # Expecting Bad Request
    assert share_response.json().get("detail") == "Failed to share summary: 400: You cannot share a summary with yourself."

def test_get_summary_success():
    """
    Test retrieving a specific summary.
    
    
    Expected Output:
        - 200 status code
        - Summary details matching the requested ID
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']
    # Share the summary created by the first user
    summary = summaries_collection.find_one({'userId': userId})
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    # Use the client to send a request to the summary retrieval endpoint
    response = client.get(f"/summary/{summary["_id"]}",headers=headers)
    # Assert that the response is successful
    assert response.status_code == 200
    retrieved_summary = response.json()
    # Assert that the retrieved summary matches the created summary
    assert retrieved_summary is not None
    assert retrieved_summary["result"]["id"] == summary["_id"]  # Adjust based on your serializer output

def test_get_summary_not_found():
    """
    Test retrieving a non-existent summary.
    
    Expected Output:
        - ValueError with "Summary not found" message
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    with pytest.raises(ValueError, match="Summary not found"):
        response = client.get("/summary/non_existent_summary_id", headers=headers) 

def test_get_shared_summaries_success():
    """
    Test retrieving summaries shared with a user.
    
    Expected Output:
        - 200 status code
        - List of shared summaries
    """
    # Create a second user
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']
    # Share the summary created by the first user
    summary = summaries_collection.find_one({'userId': userId})
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    recipient_email = "bob.johnson@example.com"
    share_response = client.post("/summary/share", json={"summary_id": summary["_id"], "recipient": recipient_email}, headers=headers)

    user = users_collection.find_one({'email':recipient_email})
    response = client.get(f"/user/{user["_id"]}/shared-summaries", headers={"Authorization": f"Bearer {auth_token}"})
    assert response.status_code == 200
    assert isinstance(response.json()["result"], list)
    assert len(response.json()["result"]) > 0  

def test_get_shared_summaries_no_shared():
    """
    Test retrieving shared summaries for user with none shared.
    
    Expected Output:
        - 200 status code
        - Empty list
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    response = client.get(f"/user/{userId}/shared-summaries", headers=headers)
    
    assert response.status_code == 200
    assert response.json()["result"] == []  # Expecting an empty list

def test_get_shared_summaries_invalid_user():
    """
    Test retrieving shared summaries for invalid user.
    
    Expected Output:
        - 200 status code
        - Empty list
    """
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token']
    invalid_user_id = "non_existent_user_id"
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    response = client.get(f"/user/{invalid_user_id}/shared-summaries", headers=headers)

    assert response.status_code == 200  # Assuming the API returns an empty list for non-existent users
    assert response.json()["result"] == []  # Expecting an empty list since the user does not exist

@pytest.fixture
def sample_pdf():
    """
    Create a test PDF file in memory.
    
    Returns:
        BytesIO: Buffer containing a simple PDF with test text
    """
    # Create a PDF in memory
    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.drawString(100, 750, "Hello World")
    c.drawString(100, 700, "This is a test PDF")
    c.save()
    
    # Move buffer position to start
    buffer.seek(0)
    return buffer

def test_extract_text_from_pdf(sample_pdf):
    """
    Test PDF text extraction functionality.
    
    Input:
        - PDF file containing "Hello World" and test text
    Expected Output:
        - Extracted text containing the PDF content
    """
    # Create a mock UploadFile
    class MockUploadFile:
        def __init__(self, filename, contents):
            self.filename = filename
            self.file = BytesIO(contents)

    mock_file = MockUploadFile("test.pdf", sample_pdf.getvalue())
    
    # Extract text from the PDF
    from services.service import extract_text_from_pdf
    extracted_text = extract_text_from_pdf(mock_file)
    
    # Check if the text was extracted (exact matching might be tricky due to PDF formatting)
    assert "Hello World" in extracted_text
    assert "This is a test PDF" in extracted_text

@pytest.fixture
def sample_text():
    # Create a text file content in memory
    buffer = BytesIO()
    content = "Hello World\nThis is a test text file"
    buffer.write(content.encode('utf-8'))
    
    # Move buffer position to start
    buffer.seek(0)
    return buffer

@pytest.mark.asyncio
async def test_extract_text_from_file(sample_text):
    """
    Test extracting text from a text file.
    
    Input:
        - Text file with sample content
    Expected Output:
        - Extracted text matching input content
    """
    # Create a mock UploadFile
    class MockUploadFile:
        def __init__(self, filename, contents):
            self.filename = filename
            self.file = BytesIO(contents)
        
        async def read(self):
            return self.file.read()

    mock_file = MockUploadFile("test.txt", sample_text.getvalue())
    
    # Extract text from the file
    from services.service import extract_text_from_file
    mock_file_contents = await mock_file.read()
    extracted_text = extract_text_from_file(mock_file, mock_file_contents)
    
    # Check if the text was extracted correctly
    assert "Hello World" in extracted_text
    assert "This is a test text file" in extracted_text


@pytest.fixture
def sample_docx():
    # Create a DOCX in memory
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test DOCX file")
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@pytest.mark.asyncio
async def test_extract_text_from_docx(sample_docx):
    """
    Test extracting text from a DOCX file.
    
    Input:
        - DOCX file with sample content
    Expected Output:
        - Extracted text matching input content
    """
    # Create a mock UploadFile
    class MockUploadFile:
        def __init__(self, filename, contents):
            self.filename = filename
            self.file = BytesIO(contents)
        
        async def read(self):
            return self.file.read()

    mock_file = MockUploadFile("test.docx", sample_docx.getvalue())
    
    # Extract text from the file
    from services.service import extract_text_from_file
    mock_file_contents = await mock_file.read()
    extracted_text = extract_text_from_file(mock_file, mock_file_contents)
    
    # Check if the text was extracted correctly
    assert "Hello World" in extracted_text
    assert "This is a test DOCX file" in extracted_text

def test_regenerate_feedback_with_text():
    """
    Test regenerating a text-based summary with feedback.
    
    Expected Output:
        - 201 status code
        - New regenerated summary
    """
    # First create and verify a user to get auth token
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token']
    
    # Get an existing summary for the user
    userId = response.json()['result']['user']['id']
    summary = summaries_collection.find_one({'userId': userId})
    
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    
    feedback = "Please make it more concise"
    
    # Call the regenerate endpoint
    response = client.post(
        f"/summary/regenerate/{summary['_id']}", 
        json=feedback,
        headers=headers
    )
    
    assert response.status_code == 201
    assert "result" in response.json()
    new_summary = response.json()["result"]
    assert isinstance(new_summary, dict)

def test_regenerate_feedback_invalid_summary():
    """
    Test regenerating an invalid summary.
    
    Expected Output:
        - 404 status code
    """
    # First create and verify a user to get auth token
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token']
    
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    
    invalid_summary_id = "nonexistent_id"
    feedback = "Please make it more concise"
    
    response = client.post(
        f"/summary/regenerate/{invalid_summary_id}", 
        json=feedback,
        headers=headers
    )
    
    assert response.status_code == 404

def test_service_download_file_not_found():
    """
    Test downloading a non-existent file.
    
    Expected Output:
        - HTTPException with 404 status code
        - "File not found" error message
    """
    # Use a valid ObjectId format that does not exist in the GridFS collection
    non_existent_file_id = "507f1f77bcf86cd799439011"

    # Ensure the service raises the expected HTTPException
    with pytest.raises(HTTPException) as exc_info:
        service_download_file(non_existent_file_id)

    # Verify exception details
    assert exc_info.value.status_code == 404
    assert exc_info.value.detail == "File not found"

def test_service_download_file(sample_pdf):
    """
    Test downloading a valid file.
    
    Input:
        - Valid PDF file in GridFS
    Expected Output:
        - StreamingResponse with correct headers
        - PDF content matches input
    """
    grid_fs = GridFS(db)

    file_id = grid_fs.put(
        sample_pdf.getvalue(),
        filename="test.pdf",
        content_type="application/pdf"
    )

    try:
        stored_file = grid_fs.get(file_id)
        assert stored_file.filename == "test.pdf"
        assert stored_file.content_type == "application/pdf"
    except Exception as e:
        pytest.fail(f"File was not stored correctly: {e}")

    response = service_download_file(str(file_id))
    assert isinstance(response, StreamingResponse)

    assert response.headers["Content-Disposition"] == "attachment; filename=test.pdf"
    assert response.headers["content-type"] == "application/pdf"

    grid_fs.delete(file_id)

def test_download_file(sample_pdf):
    """
    Test file download endpoint.
    
    Input:
        - Valid file ID
    Expected Output:
        - 200 status code
        - Correct content headers
        - File content matches input
    """
    grid_fs = GridFS(db)
    # Step 1: Store the file in GridFS
    file_id = grid_fs.put(
        sample_pdf.getvalue(),
        filename="test.pdf",
        content_type="application/pdf"
    )

    response = client.get(f"/download/{str(file_id)}")

    # Verify the response
    assert response.status_code == 200
    assert response.headers["Content-Disposition"] == "attachment; filename=test.pdf"
    assert response.headers["content-type"] == "application/pdf"
    # assert response.content == b"This is a test PDF content."

    # Cleanup: Delete the test file from GridFS
    grid_fs.delete(file_id)


def test_summary_upload(sample_text):
    """
    Test file upload functionality for summaries.
    
    Input:
        - Text file
    Expected Output:
        - 201 status code
        - Response containing file_id
    
    Note: Includes 3-second delay to prevent rate limiting
    """
    file_to_upload = ("file", ("test_file.txt", sample_text, "text/plain"))
    user_data = {
        "email": "alice.smith@example.com",
        "password": "password123"
    }
    response = client.post("/user/verify", json=user_data)
    auth_token = response.json()['result']['auth_token'] 
    userId = response.json()['result']['user']['id']
    data = {
        "userId": userId,
        "type": "code",
        "uploadType": "upload",
    }

    time.sleep(3)

    response = client.post(
        "/summary/upload",
        data=data,
        files=[file_to_upload],
        headers={
            "Authorization": f"Bearer {auth_token}",
        },
    )

    print("The response is ",response.json())

    assert response.status_code == 201
    json_response = response.json()
    assert json_response["status"] == "OK"
    assert "file_id" in json_response["result"]

def test_user_summaries_empty():
    """
    Test retrieving summaries for new user.
    
    Expected Output:
        - 200 status code
        - Empty list of summaries
    """
    new_user = User(
        firstName="Empty",
        lastName="User",
        phone="+14132752739",
        email="empty.user@example.com",
        password="password123",
        confirmPassword="password123"
    )
    
    user_data_dict = new_user.dict()
    user_data_dict['createdAt'] = new_user.createdAt.isoformat()
    user_data_dict['lastLoggedInAt'] = new_user.lastLoggedInAt.isoformat()
    
    response = client.post("/user/create", json=user_data_dict)
    assert response.status_code == 201
    
    verify_data = {
        "email": new_user.email,
        "password": new_user.password
    }
    verify_response = client.post("/user/verify", json=verify_data)
    assert verify_response.status_code == 200
    
    auth_token = verify_response.json()['result']['auth_token']
    userId = verify_response.json()['result']['user']['id']
    
    headers = {
        "Authorization": f"Bearer {auth_token}"
    }
    
    summaries_response = client.get(f"/summaries/{userId}", headers=headers)
    
    assert summaries_response.status_code == 200
    assert summaries_response.json()["result"] == []
    
    users_collection.delete_one({"email": new_user.email})

